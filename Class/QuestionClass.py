from Class.CourseClass import course
from Class.TestClass import test
import PyPDF3
import pdf2image
import docx
import docx.shared


class question:
    def __init__(self, course, test, qc, mt, st, dif, sult, Format, te):
        self.course = course
        self.test = test
        self.question_code = qc
        self.main_topic = mt
        self.sub_theme = st
        self.difficulty = dif
        self.solution_type = sult
        self.format = Format
        self.test_exams = te
        self.path = "pdfFileHere\{0}.pdf".format(qc)
        test.__crop__()
        self.pdf = PyPDF3.PdfFileReader(self.path, "rb")
        self.image = pdf2image.convert_from_path(self.path)

    def imageToWord(self):
        docs=docx.Document()
        parg=docs.add_paragraph()
        r=p.add_run()
        r.add_text('question number 1:\n')
        r.add_picture('exemple.png')
        docs.save('exe.docx')

    def __str__(self):
        return '{0} {1} {2} {3} {4} {5} {6} {7} {8} {9} {10} {11} {12} {13}'.format(self.course.name_course,
                                                                                    self.course.year,
                                                                                    self.course.semester,
                                                                                    self.test.year, self.test.semester,
                                                                                    self.test.deadline,
                                                                                    self.question_code, self.main_topic,
                                                                                    self.sub_theme, self.difficulty,
                                                                                    self.solution_type, self.format,
                                                                                    self.test_exams, self.path)
path=input("path\n")
test1=test(1,1,1,path)
course1=course(1,1,1)
q1=question(course1,test1,112,1,1,1,1,1,1)
q1.imageToWord()